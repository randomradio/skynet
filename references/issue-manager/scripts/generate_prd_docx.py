#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成产品需求文档 PRD 的 Word 版本，含项目代码功能对应说明
输出到 Download 目录
"""
import os
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("请先安装 python-docx: pip install python-docx")
    exit(1)


def set_cell_border(cell, **kwargs):
    """设置表格单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        edge_attr = kwargs.get(edge)
        if edge_attr:
            tag = 'w:{}'.format(edge)
            element = OxmlElement(tag)
            element.set(qn('w:val'), edge_attr.get('val', 'single'))
            element.set(qn('w:sz'), str(edge_attr.get('sz', 4)))
            element.set(qn('w:color'), edge_attr.get('color', '000000'))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16 if level == 1 else 14 if level == 2 else 12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    p.space_before = Pt(12)
    p.space_after = Pt(6)
    return p


def add_para(doc, text, bold=False):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
        run.font.size = Pt(10.5)
        if bold:
            run.bold = True
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    hrow = table.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        for p in hrow.cells[i].paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.name = 'Microsoft YaHei'
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        r = table.rows[ri + 1]
        for ci, cell in enumerate(row):
            if ci < len(r.cells):
                r.cells[ci].text = str(cell)
                for p in r.cells[ci].paragraphs:
                    for r0 in p.runs:
                        r0.font.name = 'Microsoft YaHei'
                        r0.font.size = Pt(9)
    return table


def main():
    doc = Document()
    doc.styles['Normal'].font.name = 'Microsoft YaHei'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    doc.styles['Normal'].font.size = Pt(10.5)

    # 标题
    title = doc.add_paragraph()
    t = title.add_run('GitHub Issue 智能管理系统 - 产品需求文档（PRD）')
    t.bold = True
    t.font.size = Pt(18)
    t.font.name = 'Microsoft YaHei'
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.space_after = Pt(6)
    add_para(doc, '版本 V1.0 | 2026-02-24')

    # 一、产品是什么
    add_heading(doc, '一、产品是什么', 1)
    add_para(doc, '自动从 GitHub 拉取 Issue，用 AI 解析类型/优先级/关联关系，存库后做日报、多维分析、AI 洞察，并通过邮件推送。支持 matrixone、matrixflow 两仓库（邮件中 matrixflow 优先）。')

    # 二、解决什么问题
    add_heading(doc, '二、解决什么问题', 1)
    add_para(doc, '• 缺乏全局视角 — 多仓 Issue 难统一看进展')
    add_para(doc, '• 更新追踪难 — 变化靠人工，无历史记录')
    add_para(doc, '• 关联不清晰 — 依赖/阻塞/共用 Feature 难识别')
    add_para(doc, '• 缺乏智能化 — 需要自动分析、自动报告')

    # 三、主要功能
    add_heading(doc, '三、主要功能', 1)
    add_table(doc, ['功能', '说明'], [
        ['数据采集', '全量/增量同步 Issue、评论、时间线'],
        ['智能解析', 'AI 提取类型、优先级、模块、关联关系'],
        ['日报与进度', '每日统计、按状态/类型/优先级汇总'],
        ['可扩展分析', '标签、模块、客户、趋势等，配置驱动'],
        ['多维度分析', '客户项目、共用 Feature、阻塞链'],
        ['AI 驱动分析', '按 L1→L2→L3→L4 输出健康度与建议（matrixflow+matrixone）'],
        ['邮件推送', '日报 + 可扩展分析 + AI 分析 写入邮件'],
        ['关联补全', '脚本补全历史遗漏的关联关系'],
    ])

    # 四、怎么用
    add_heading(doc, '四、怎么用', 1)
    add_para(doc, '推荐用 auto_run.py（可配置、可定时）：')
    add_para(doc, '• 完整流程（同步+报告+邮件）：')
    add_para(doc, '  python3 auto_run.py --repo-owner matrixorigin --repo-name matrixflow --email 收件人@example.com')
    add_para(doc, '• 只做分析并发邮件（不同步）：')
    add_para(doc, '  python3 auto_run.py --repo-owner matrixorigin --repo-name matrixflow --skip-sync --extensible-report --ai-report --email 收件人@example.com')
    add_para(doc, '• 补全关联：加 --supplement-relations；多维度报告：加 --multi-report')
    add_para(doc, '交互式运行：python3 main.py 或 python3 run.py。脚本在 scripts/ 下，如 supplement_relations.py、run_ai_analysis.py 等。')

    # 五、功能与代码对应
    add_heading(doc, '五、功能与代码对应', 1)
    add_para(doc, '让未用过项目的人快速知道「功能在哪儿实现」。')
    add_table(doc, ['功能', '代码位置'], [
        ['数据采集', 'modules/github_collector/；main.py 调用'],
        ['智能解析', 'modules/llm_parser/'],
        ['存储', 'modules/database_storage/mo_client.py'],
        ['日报与进度', 'modules/analysis_engine/'],
        ['可扩展分析', 'modules/analysis_extensible/；config/analysis_config.yaml'],
        ['多维度分析', 'modules/analysis_engine/multi_dimensional_analyzer.py'],
        ['AI 驱动分析', 'modules/ai_analysis/ai_driven_analysis_engine.py'],
        ['邮件', 'modules/email_sender/email_sender.py'],
        ['关联补全', 'main.py Phase2 + scripts/supplement_relations.py'],
    ])
    add_para(doc, '入口：main.py / run.py（交互）；auto_run.py（推荐）。配置：config/config.py。详细见 README、docs/。')

    add_para(doc, '')
    add_para(doc, '—— 文档结束 ——')

    # 保存
    out_dir = Path.home() / 'Downloads'
    out_path = out_dir / 'GitHub_Issue智能管理系统_产品需求文档PRD-V1-260224.docx'
    doc.save(str(out_path))
    print(f"已生成: {out_path}")
    return str(out_path)


if __name__ == '__main__':
    main()
